"""
Genera el Manual de Usuario de Digital One en formato Word (.docx)
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Rutas ──
BASE = r"C:\Users\Gustavo\OneDrive\Gustavo\Empresas\IntegraIA\WebHTML"
IMG_DIR = os.path.join(BASE, "images", "digital-one")
OUTPUT = os.path.join(BASE, "Docs", "Manual_Usuario_DigitalOne.docx")

# ── Colores ──
DARK_BLUE = RGBColor(0x0A, 0x16, 0x28)
ACCENT_TEAL = RGBColor(0x00, 0xC9, 0xA7)
ACCENT_BLUE = RGBColor(0x00, 0x80, 0xFF)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
TEXT_DARK = RGBColor(0x1A, 0x2A, 0x4A)
TEXT_GRAY = RGBColor(0x4A, 0x4A, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF0, 0xF4, 0xFA)

doc = Document()

# ── Estilos base ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = TEXT_DARK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = DARK_BLUE
            run.font.size = Pt(26)
            run.font.name = 'Calibri Light'
        elif level == 2:
            run.font.color.rgb = ACCENT_TEAL
            run.font.size = Pt(18)
            run.font.name = 'Calibri'
        elif level == 3:
            run.font.color.rgb = TEXT_DARK
            run.font.size = Pt(14)
            run.font.name = 'Calibri'
    return h


def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5E9" w:val="clear"/>')
    pPr.append(shading)
    run_icon = p.add_run("💡 ")
    run_icon.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    run.italic = True
    return p


def add_important(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3E0" w:val="clear"/>')
    pPr.append(shading)
    run_icon = p.add_run("⚠️ ")
    run_icon.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
    run.bold = True
    return p


def add_image(filename, width=Inches(5.5), caption=None):
    img_path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(img_path):
        add_body(f"[Imagen: {filename}]", italic=True, color=TEXT_GRAY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = TEXT_GRAY
        r.italic = True
        cap.paragraph_format.space_after = Pt(12)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = WHITE
                run.font.name = 'Calibri'
        # Dark background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0A1628" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            # Alternate row shading
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()  # spacing
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p


def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)


# ══════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════

# Espacios antes del título
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Digital One")
run.font.size = Pt(48)
run.font.name = 'Calibri Light'
run.font.color.rgb = DARK_BLUE
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Manual del Usuario")
run2.font.size = Pt(28)
run2.font.name = 'Calibri Light'
run2.font.color.rgb = ACCENT_TEAL

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Sistema de Control de Asistencia y Gestión de Personal")
run3.font.size = Pt(14)
run3.font.color.rgb = TEXT_GRAY
run3.font.name = 'Calibri'

for _ in range(4):
    doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("Versión 7.0")
run4.font.size = Pt(12)
run4.font.color.rgb = TEXT_GRAY
run4.font.name = 'Calibri'

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("Marzo 2026")
run5.font.size = Pt(12)
run5.font.color.rgb = TEXT_GRAY

doc.add_paragraph()
doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run6 = p6.add_run("Integra IA SRL")
run6.font.size = Pt(14)
run6.font.color.rgb = GOLD
run6.bold = True
run6.font.name = 'Calibri'

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run7 = p7.add_run("www.integraia.tech")
run7.font.size = Pt(11)
run7.font.color.rgb = ACCENT_BLUE

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════

add_heading_styled("Índice", level=1)
doc.add_paragraph()

toc_items = [
    ("1.", "Introducción"),
    ("2.", "Instalación"),
    ("   2.1", "Instalador Completo (Local)"),
    ("   2.2", "Instalador Liviano (Nube)"),
    ("3.", "Primeros Pasos"),
    ("4.", "Digital One Fichador"),
    ("   4.1", "Fichada por Huella Digital"),
    ("   4.2", "Fichada por PIN"),
    ("   4.3", "Fichada Móvil"),
    ("   4.4", "Detección automática de modo"),
    ("5.", "Digital One Administrador"),
    ("   5.1", "Gestión de Legajos"),
    ("   5.2", "Gestión de Fichadas"),
    ("   5.3", "Tablas del Sistema"),
    ("   5.4", "Configuración"),
    ("   5.5", "Reportes"),
    ("6.", "Portal Web"),
    ("7.", "Fichada Móvil — Digital One Mobile"),
    ("8.", "Sistema de Licencias"),
    ("9.", "Preguntas Frecuentes"),
    ("10.", "Soporte Técnico"),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{num}  ")
    r1.font.size = Pt(11)
    r1.font.color.rgb = ACCENT_TEAL
    r1.bold = True
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEXT_DARK
    p.paragraph_format.space_after = Pt(2)
    if not num.startswith("   "):
        p.paragraph_format.space_before = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════

add_heading_styled("1. Introducción", level=1)

add_heading_styled("¿Qué es Digital One?", level=2)

add_body("Digital One es un sistema de control de asistencia y gestión de personal que permite registrar los ingresos y egresos del personal de su empresa. El sistema ofrece múltiples modalidades de fichada para adaptarse a las necesidades de cada organización:")

add_bullet("El empleado apoya su dedo en un lector USB y el sistema lo identifica automáticamente.", bold_prefix="Huella digital: ")
add_bullet("El empleado ingresa su número de legajo y un código PIN personal.", bold_prefix="PIN: ")
add_bullet("El empleado ficha desde su smartphone accediendo a la PWA del portal, con validación de ubicación GPS.", bold_prefix="Móvil: ")
add_bullet("Para pruebas sin necesidad de hardware.", bold_prefix="Modo demostración: ")

add_image("hero-digital-one.png", width=Inches(5.5), caption="Digital One — Visión general del sistema")

add_heading_styled("Componentes del sistema", level=2)

add_body("Digital One se compone de cuatro aplicaciones que trabajan juntas:")

add_table(
    ["Aplicación", "Función"],
    [
        ["Digital One Fichador", "Terminal de fichaje donde los empleados registran entrada y salida"],
        ["Digital One Administrador", "Aplicación de gestión: legajos, horarios, reportes, configuración"],
        ["Portal Web", "Acceso vía navegador para consulta de fichadas, reportes y gestión remota"],
        ["Digital One Mobile", "PWA de smartphone para fichada con validación de ubicación GPS"],
    ]
)

add_heading_styled("Requisitos del equipo", level=2)

add_table(
    ["Requisito", "Detalle"],
    [
        ["Sistema operativo", "Windows 10 o Windows 11 (64 bits)"],
        [".NET Framework", "4.8 o superior (normalmente ya incluido en Windows 10/11)"],
        ["Permisos", "Administrador local para instalar"],
        ["Lector de huellas", "DigitalPersona uAreU 4500 (opcional si usa PIN o móvil)"],
        ["Conexión a internet", "Requerida para instalación en nube y fichada móvil; opcional para local"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. INSTALACIÓN
# ══════════════════════════════════════════════════════════════

add_heading_styled("2. Instalación", level=1)

add_body("Digital One cuenta con dos instaladores según el tipo de despliegue:")

add_table(
    ["Instalador", "Tamaño", "Para quién"],
    [
        ["Instalador Completo", "~180 MB", "Empresas que quieren su propia base de datos local"],
        ["Instalador Liviano (Cloud)", "~25 MB", "Empresas que usan la base de datos en la nube"],
    ]
)

add_note("Ambos instaladores instalan Fichador y Administrador juntos. La diferencia es dónde se ubica la base de datos.")

add_separator()

# 2.1
add_heading_styled("2.1 Instalador Completo (Local)", level=2)

add_body("Este instalador crea todo lo necesario en su propio equipo, incluyendo SQL Server Express si no lo tiene instalado.")

add_heading_styled("Paso a paso", level=3)

add_numbered("Ejecute el instalador: haga doble clic en DigitalPlus_Suite_Setup_v1.3.exe")
add_numbered("En la pantalla de bienvenida, haga clic en Siguiente.")
add_numbered("Seleccione el modo de instalación Local y haga clic en Siguiente.")
add_numbered("Si su equipo no tiene SQL Server Express, el instalador lo descargará e instalará automáticamente. Este proceso puede tomar entre 5 y 15 minutos.")
add_numbered("El instalador creará la base de datos con todas las tablas necesarias.")
add_numbered("Seleccione los accesos directos deseados (Fichador y Administrador en el escritorio, inicio automático con Windows).")
add_numbered("Haga clic en Instalar y espere a que finalice.")
add_numbered("Haga clic en Finalizar.")

add_note("Si Windows muestra una advertencia de seguridad (SmartScreen), haga clic en \"Más información\" y luego en \"Ejecutar de todas formas\".")

add_separator()

# 2.2
add_heading_styled("2.2 Instalador Liviano (Nube)", level=2)

add_body("Este es el instalador recomendado para la mayoría de las empresas. La base de datos está alojada en la nube, es rápido y liviano (~25 MB), pero requiere un código de activación proporcionado por el administrador del sistema.")

add_heading_styled("Qué necesita antes de empezar", level=3)

add_bullet("El archivo del instalador: DigitalPlus_Cloud_Setup_v1.0.exe")
add_bullet("El código de activación de su empresa (proporcionado por su proveedor)")
add_bullet("Conexión a internet (obligatoria durante la instalación)")

add_heading_styled("Paso a paso", level=3)

add_numbered("Ejecute el instalador: haga doble clic en DigitalPlus_Cloud_Setup_v1.0.exe")
add_numbered("En la pantalla de bienvenida, haga clic en Siguiente.")
add_numbered("Seleccione los accesos directos y opciones deseadas.")
add_numbered("Ingrese el código de activación y haga clic en Validar Código. Espere la respuesta del servidor:")
add_bullet("Verde: Código válido. Puede continuar.", bold_prefix="✅ ")
add_bullet("Rojo: Código inválido o expirado. Verifique el código.", bold_prefix="❌ ")
add_numbered("Haga clic en Instalar. El instalador realizará automáticamente la copia de archivos, instalación del driver del lector, configuración de conexión a la nube y cifrado de seguridad.")
add_numbered("Haga clic en Finalizar.")

add_important("Sin un código válido no podrá continuar con la instalación. Cada código está asociado a una empresa específica.")

add_heading_styled("Auto-registro de terminales", level=3)

add_body("Al abrir el Fichador por primera vez, el sistema realiza un auto-registro de la terminal: detecta el nombre de la computadora y la registra automáticamente en la base de datos, asociándola a la sucursal principal de su empresa.")

add_note("No es necesario dar de alta la computadora manualmente desde el portal web ni desde el Administrador.")

add_separator()

add_heading_styled("Conectar el lector de huellas", level=2)

add_body("Después de instalar, conecte el lector DigitalPersona uAreU 4500 al puerto USB. Windows instalará los drivers automáticamente (ya fueron incluidos durante la instalación).")

add_note("Si el lector ya estaba conectado antes de instalar, desconéctelo y vuelva a conectarlo.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. PRIMEROS PASOS
# ══════════════════════════════════════════════════════════════

add_heading_styled("3. Primeros Pasos", level=1)

add_body("Una vez instalado Digital One, siga este orden recomendado de configuración:")

add_numbered("Abra Digital One Administrador")
add_numbered("Configure los datos básicos: Sucursales, Categorías, Horarios, Sectores")
add_numbered("Cargue los Legajos (empleados) — puede tomarles foto con la cámara web")
add_numbered("Registre las huellas digitales de cada empleado (o asigne PINs)")
add_numbered("Abra Digital One Fichador en la terminal de fichaje — la computadora se registra automáticamente")
add_numbered("Pruebe fichando con un empleado registrado")

add_note("No es necesario registrar las terminales (computadoras de fichaje) manualmente. El Fichador las registra automáticamente la primera vez que se ejecuta en cada computadora.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. DIGITAL ONE FICHADOR
# ══════════════════════════════════════════════════════════════

add_heading_styled("4. Digital One Fichador", level=1)

add_body("El Fichador es la aplicación que los empleados usan día a día para registrar su entrada y salida. Al abrirlo, verá la pantalla de fichaje con el nombre de la sucursal, el reloj con fecha y hora actual, el área de semáforo visual y el panel de fichada.")

add_image("fichador-semaforo.png", width=Inches(5), caption="Pantalla principal del Fichador con semáforo visual")

add_heading_styled("Semáforo visual", level=3)
add_body("Al fichar, el sistema muestra un semáforo de colores que confirma visualmente el resultado:")
add_bullet("Fichada registrada correctamente. Muestra nombre del empleado y si es ENTRADA o SALIDA.", bold_prefix="🟢 Verde: ")
add_bullet("Procesando captura de huella.", bold_prefix="🟡 Amarillo: ")
add_bullet("Huella no reconocida. El empleado debe reintentar.", bold_prefix="🔴 Rojo: ")

add_separator()

# 4.1
add_heading_styled("4.1 Fichada por Huella Digital", level=2)

add_body("Este es el modo principal si tiene un lector de huellas conectado. El proceso es simple:")

add_numbered("El empleado apoya su dedo en el lector")
add_numbered("El semáforo cambia a amarillo (procesando)")
add_numbered("Resultado: verde si la fichada fue exitosa, rojo si la huella no fue reconocida")

add_image("fichador-huella.png", width=Inches(5), caption="Fichador — Modo Huella Digital")

add_separator()

# 4.2
add_heading_styled("4.2 Fichada por PIN", level=2)

add_body("Si no hay lector de huellas o si el modo PIN está habilitado, los empleados pueden fichar ingresando su número de legajo y PIN personal:")

add_numbered("El empleado ingresa su número de legajo")
add_numbered("Ingresa su PIN de 4 a 6 dígitos")
add_numbered("El sistema valida y registra la fichada")

add_image("fichador-pin.png", width=Inches(5), caption="Fichador — Modo PIN")

add_heading_styled("Primera vez (sin PIN asignado)", level=3)
add_body("Si el legajo no tiene un PIN asignado, el sistema le preguntará \"¿Desea crear uno ahora?\" con opciones Sí y No. Si acepta, deberá ingresar un nuevo PIN (4 a 6 dígitos) y confirmarlo.")

add_heading_styled("Cambio de PIN", level=3)
add_body("Existen dos formas de cambiar el PIN:")
add_bullet("En la pantalla de fichada hay un link \"Cambiar mi PIN\" donde el empleado ingresa su legajo, PIN actual, nuevo PIN y confirmación.", bold_prefix="Cambio voluntario: ")
add_bullet("Cuando el administrador marca \"Forzar cambio\" o \"Resetear PIN\", la próxima vez que el empleado ingrese su legajo, deberá crear un nuevo PIN.", bold_prefix="Cambio forzado: ")

add_heading_styled("PIN expirado", level=3)
add_body("Si el PIN del empleado está vencido, al ingresar el legajo y PIN actual correctamente, el sistema le solicitará que cree un nuevo PIN antes de completar la fichada.")

add_separator()

# 4.3
add_heading_styled("4.3 Fichada Móvil", level=2)

add_body("Con Digital One Mobile, los empleados pueden fichar directamente desde su smartphone. No necesitan instalar ninguna aplicación: acceden a la PWA (Progressive Web App) desde el navegador del celular.")

add_image("fichador-movil.png", width=Inches(4), caption="Fichador — Modo Móvil desde smartphone")

add_body("El proceso de fichada móvil es el siguiente:")
add_numbered("El empleado abre el navegador de su celular y accede a la PWA de Digital One")
add_numbered("Se identifica con su número de legajo y PIN")
add_numbered("El sistema valida automáticamente su ubicación por GPS")
add_numbered("Si está dentro del radio autorizado de la sucursal, se habilita el botón de fichada")
add_numbered("El empleado presiona \"Fichar\" y la entrada o salida queda registrada")

add_note("La fichada móvil funciona en cualquier smartphone con Android o iOS. Solo requiere un navegador moderno, conexión a internet y GPS habilitado.")

add_important("El administrador debe configurar la validación de ubicación GPS/WiFi por sucursal desde el Portal Web antes de que los empleados puedan usar la fichada móvil.")

add_separator()

# 4.4
add_heading_styled("4.4 Detección automática de modo", level=2)

add_body("El Fichador detecta automáticamente si hay un lector de huellas conectado:")
add_bullet("Usa modo Huella", bold_prefix="Lector detectado: ")
add_bullet("Cambia automáticamente a modo PIN", bold_prefix="Sin lector + PIN habilitado: ")
add_bullet("Cambia a modo Demo", bold_prefix="Sin lector + Demo habilitado: ")

add_note("Si se desconecta el lector durante el uso, el Fichador cambia de modo en tiempo real.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. DIGITAL ONE ADMINISTRADOR
# ══════════════════════════════════════════════════════════════

add_heading_styled("5. Digital One Administrador", level=1)

add_body("La aplicación Administrador permite gestionar todos los aspectos del control de asistencia: empleados, horarios, sucursales, reportes y más. Al abrirla, verá un menú lateral con todas las opciones disponibles, el logo de su empresa y accesos a redes sociales.")

add_image("administrador-principal.png", width=Inches(5.5), caption="Pantalla principal del Administrador")

add_separator()

# 5.1
add_heading_styled("5.1 Gestión de Legajos (Empleados)", level=2)

add_body("Desde RRHH > Legajos puede gestionar toda la información de sus empleados:")

add_bullet("Alta, baja y modificación de legajos: nombre, legajo, categoría, horario, sector, sucursal", bold_prefix="")
add_bullet("Tomar foto del empleado usando la cámara web")
add_bullet("Registrar huellas digitales (requiere lector conectado)")
add_bullet("Asignar o resetear PIN")

add_heading_styled("Registrar huellas", level=3)
add_numbered("Seleccione el empleado en la lista")
add_numbered("Haga clic en el botón de huella")
add_numbered("Siga las instrucciones del asistente de enrolamiento")
add_numbered("Apoye el mismo dedo 4 veces para capturar la huella")
add_numbered("Repita con otro dedo si lo desea")

add_image("administrador-huellas.png", width=Inches(5), caption="Registro de huellas digitales — Asistente de enrolamiento")

add_separator()

# 5.2
add_heading_styled("5.2 Gestión de Fichadas", level=2)

add_body("Desde RRHH > Fichadas puede:")
add_bullet("Ver todas las fichadas registradas")
add_bullet("Filtrar por empleado, fecha, tipo (Entrada/Salida)")
add_bullet("Agregar fichadas manuales (en caso de olvido del empleado)")
add_bullet("Ver llegadas tarde")

add_separator()

# 5.3
add_heading_styled("5.3 Tablas del Sistema", level=2)

add_table(
    ["Opción", "Qué gestiona"],
    [
        ["Sucursales", "Ubicaciones físicas de la empresa"],
        ["Categorías", "Clasificación de empleados (ej: Operario, Administrativo)"],
        ["Horarios", "Definición de horarios de trabajo"],
        ["Sectores", "Áreas de la empresa (ej: Producción, RRHH)"],
        ["Incidencias", "Tipos de ausencia (vacaciones, enfermedad, permiso)"],
        ["Feriados", "Días no laborables"],
    ]
)

add_separator()

# 5.4
add_heading_styled("5.4 Configuración del Sistema", level=2)

add_heading_styled("Pestaña Fichada", level=3)
add_bullet("Habilitar o deshabilitar fichada por PIN", bold_prefix="Modo PIN: ")
add_bullet("Cantidad de días para que el PIN expire (0 = no expira)", bold_prefix="Expiración de PIN: ")
add_bullet("Habilitar o deshabilitar modo demostración", bold_prefix="Modo Demo: ")

add_heading_styled("Pestaña PINs", level=3)
add_body("Muestra todos los legajos con su estado de PIN. Incluye filtros por:")
add_bullet("Todos los legajos")
add_bullet("Con PIN activo")
add_bullet("Sin PIN")
add_bullet("Vencidos")
add_bullet("Cambio pendiente")

add_body("Acciones disponibles:")
add_bullet("Marca los legajos seleccionados para que deban establecer un nuevo PIN en su próxima fichada.", bold_prefix="Forzar cambio de PIN: ")
add_bullet("Elimina el PIN del empleado por completo. La próxima vez se le solicitará crear uno nuevo.", bold_prefix="Resetear PIN: ")

add_separator()

# 5.5
add_heading_styled("5.5 Reportes", level=2)

add_body("El Administrador genera reportes completos que incluyen:")
add_bullet("Asistencia por período")
add_bullet("Horas trabajadas")
add_bullet("Llegadas tarde")
add_bullet("Horas extras")

add_note("Los reportes se pueden exportar a PDF y Excel directamente desde la aplicación.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. PORTAL WEB
# ══════════════════════════════════════════════════════════════

add_heading_styled("6. Portal Web", level=1)

add_body("El Portal Web le permite acceder a toda la información de asistencia desde cualquier navegador, en cualquier dispositivo. Abra su navegador y vaya a la dirección proporcionada por su administrador.")

add_heading_styled("Acceso e inicio de sesión", level=2)
add_body("Ingrese sus credenciales de acceso (usuario y contraseña). Si es la primera vez que accede con la contraseña temporal, el sistema le pedirá que cambie su contraseña obligatoriamente antes de poder usar el portal.")

add_important("Si recibe el mensaje \"El acceso a su empresa ha sido suspendido\", contacte a su proveedor.")

add_image("portal-dashboard.png", width=Inches(5.5), caption="Portal Web — Dashboard principal con métricas en tiempo real")

add_heading_styled("Funcionalidades disponibles", level=2)

add_table(
    ["Sección", "Qué puede hacer"],
    [
        ["Legajos", "Ver y gestionar datos de empleados"],
        ["Fichadas", "Consultar fichadas por empleado y período"],
        ["Horarios", "Ver y modificar horarios de trabajo"],
        ["Categorías", "Gestionar categorías de empleados"],
        ["Sectores", "Gestionar sectores y áreas"],
        ["Sucursales", "Gestionar ubicaciones"],
        ["Terminales", "Ver terminales de fichaje registradas"],
        ["Terminales Móviles", "Gestionar dispositivos móviles, generar códigos de activación"],
        ["Fichado Móvil", "Configurar validación GPS/WiFi por sucursal"],
        ["PIN Móvil", "Asignar, cambiar y resetear PIN de empleados"],
        ["Incidencias", "Cargar permisos, ausencias, vacaciones"],
        ["Feriados", "Gestionar días feriados"],
        ["Variables", "Configuración general del sistema"],
        ["Usuarios", "Gestionar accesos al portal"],
    ]
)

add_image("portal-legajos.png", width=Inches(5.5), caption="Portal Web — Gestión de empleados")

add_image("portal-fichadas.png", width=Inches(5.5), caption="Portal Web — Consulta de fichadas")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. FICHADA MÓVIL
# ══════════════════════════════════════════════════════════════

add_heading_styled("7. Fichada Móvil — Digital One Mobile", level=1)

add_body("Digital One Mobile es una PWA (Progressive Web App) que permite a los empleados fichar directamente desde su smartphone, sin necesidad de instalar ninguna aplicación desde una tienda de apps. Funciona en cualquier celular con Android o iOS.")

add_heading_styled("¿Cómo funciona?", level=2)

add_numbered("El administrador genera un código de activación móvil desde el Portal Web (sección \"Terminales Móviles\")")
add_numbered("El empleado abre el navegador de su smartphone y accede a la dirección de la PWA")
add_numbered("Activa su dispositivo con el código proporcionado")
add_numbered("A partir de ese momento, puede fichar ingresando su legajo y PIN")
add_numbered("El sistema valida su ubicación por GPS antes de permitir la fichada")

add_heading_styled("Validación de ubicación", level=2)

add_body("El sistema ofrece dos métodos de validación para garantizar que el empleado esté en la sucursal correspondiente:")

add_bullet("El sistema verifica que el empleado esté dentro del radio configurado para la sucursal. La precisión se muestra en pantalla.", bold_prefix="GPS: ")
add_bullet("Se pueden configurar redes WiFi autorizadas por sucursal como método alternativo de validación.", bold_prefix="WiFi: ")

add_note("La configuración de validación GPS/WiFi se realiza por sucursal desde el Portal Web, sección \"Fichado Móvil\".")

add_heading_styled("PIN Móvil", level=2)

add_body("Los empleados se identifican con su número de legajo y PIN personal. Las mismas políticas de seguridad y expiración aplican tanto para la fichada de escritorio como para la móvil. El administrador puede asignar, cambiar y resetear PINs desde el formulario de Legajos en el Portal Web.")

add_heading_styled("Gestión de terminales móviles", level=2)

add_body("Desde la sección \"Terminales Móviles\" del Portal Web, el administrador puede:")
add_bullet("Ver todos los dispositivos móviles registrados")
add_bullet("Generar nuevos códigos de activación")
add_bullet("Desactivar terminales móviles")
add_bullet("Ver el historial de fichadas móviles")

add_important("Cada terminal móvil queda asociada al dispositivo y al empleado. Si el empleado cambia de celular, el administrador deberá generar un nuevo código de activación.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. SISTEMA DE LICENCIAS
# ══════════════════════════════════════════════════════════════

add_heading_styled("8. Sistema de Licencias", level=1)

add_important("En instalaciones multi-tenant (nube), el sistema de licencias funciona de manera diferente. La activación se realiza durante la instalación mediante el código de activación. El sistema de trial descrito a continuación aplica únicamente a instalaciones LOCALES.")

add_heading_styled("Período de prueba (Trial)", level=2)

add_body("Al instalar Digital One por primera vez en modo local, el sistema se activa en modo Trial automáticamente:")

add_table(
    ["Limitación", "Valor"],
    [
        ["Duración", "14 días"],
        ["Cantidad máxima de empleados", "5 legajos"],
        ["Funcionalidades", "Todas disponibles"],
    ]
)

add_heading_styled("¿Qué pasa cuando vence el Trial?", level=2)

add_body("Si el período de prueba vence o supera los 5 empleados:")
add_numbered("Se muestra una pantalla de bloqueo")
add_numbered("El sistema no permite fichar hasta que se active una licencia")
add_numbered("Los datos no se pierden")

add_heading_styled("Activar una licencia", level=2)

add_numbered("Solicite un código de activación a su proveedor")
add_numbered("Ingrese el código en la pantalla de bloqueo o desde el menú Licencias")
add_numbered("El sistema se conecta al servidor y activa el plan contratado")
add_numbered("Las limitaciones se levantan según el plan")

add_heading_styled("Requisitos de conectividad", level=2)

add_body("El sistema de licencias requiere conexión a internet periódica:")
add_bullet("Al iniciar la aplicación (para validar la licencia)")
add_bullet("Cada 4 horas (heartbeat automático)")
add_bullet("Si pierde conectividad por más de 72 horas, el sistema se bloqueará temporalmente hasta recuperar la conexión")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. PREGUNTAS FRECUENTES
# ══════════════════════════════════════════════════════════════

add_heading_styled("9. Preguntas Frecuentes", level=1)

faqs = [
    ("El Fichador no reconoce el lector de huellas",
     "1) Desconecte y vuelva a conectar el lector USB. 2) Espere 10 segundos a que Windows lo reconozca. 3) Si persiste, reinicie la aplicación. 4) Si el lector no funciona, puede usar modo PIN como alternativa."),
    ("Un empleado olvidó fichar",
     "El Administrador puede cargar fichadas manuales desde RRHH > Fichadas > Fichada Manual."),
    ("Se perdió el código de activación",
     "Contacte a su proveedor (Integra IA) para que le regenere un nuevo código de activación."),
    ("El sistema dice \"Licencia bloqueada por falta de conexión\"",
     "Verifique que el equipo tenga conexión a internet. Una vez restaurada la conexión, reinicie la aplicación y se desbloqueará automáticamente."),
    ("¿Puedo instalar Fichador en varias máquinas?",
     "Sí. Cada terminal de fichaje necesita su propia instalación. Todas las terminales pueden apuntar a la misma base de datos."),
    ("¿Cómo desinstalo Digital One?",
     "Vaya a Panel de Control > Programas > Desinstalar un programa, seleccione Digital One y siga las instrucciones. La desinstalación no elimina la base de datos ni los registros de fichadas."),
    ("¿La fichada móvil funciona sin internet?",
     "No. La fichada móvil requiere conexión a internet activa y GPS habilitado para validar la ubicación del empleado."),
    ("¿Puedo usar la fichada móvil en cualquier celular?",
     "Sí. Digital One Mobile es una PWA que funciona en cualquier smartphone con un navegador moderno (Chrome, Safari, Firefox), tanto en Android como en iOS."),
]

for question, answer in faqs:
    add_heading_styled(question, level=3)
    add_body(answer)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. SOPORTE TÉCNICO
# ══════════════════════════════════════════════════════════════

add_heading_styled("10. Soporte Técnico", level=1)

add_body("Ante cualquier inconveniente, contacte al administrador del sistema con la siguiente información:")

add_bullet("Mensaje de error exacto (captura de pantalla si es posible)")
add_bullet("Nombre del equipo")
add_bullet("Qué estaba haciendo cuando ocurrió el error")
add_bullet("Sistema operativo instalado")

doc.add_paragraph()

add_heading_styled("Contacto — Integra IA", level=2)

p = doc.add_paragraph()
r1 = p.add_run("Sitio web: ")
r1.font.size = Pt(12)
r1.bold = True
r2 = p.add_run("www.integraia.tech")
r2.font.size = Pt(12)
r2.font.color.rgb = ACCENT_BLUE

p2 = doc.add_paragraph()
r3 = p2.add_run("Email: ")
r3.font.size = Pt(12)
r3.bold = True
r4 = p2.add_run("soporte@digitaloneplus.com")
r4.font.size = Pt(12)
r4.font.color.rgb = ACCENT_BLUE

# Espaciado final
for _ in range(4):
    doc.add_paragraph()

add_separator()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_final = p_final.add_run("Fin del Manual del Usuario — Digital One v7.0")
r_final.font.size = Pt(10)
r_final.font.color.rgb = TEXT_GRAY
r_final.italic = True

p_brand = doc.add_paragraph()
p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_brand = p_brand.add_run("© 2026 Integra IA SRL — Todos los derechos reservados")
r_brand.font.size = Pt(9)
r_brand.font.color.rgb = TEXT_GRAY

# ── GUARDAR ──
doc.save(OUTPUT)
print(f"Manual generado exitosamente en:\n{OUTPUT}")
